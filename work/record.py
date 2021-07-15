import sys
import cv2
import time
import numpy as np
import pickle
import compress_pickle
from PIL import ImageGrab
from board import Board
import random
from my_keyboard import MyKeyboard
import multiprocessing
from data_manager import DataManager
from train import History
import pygame
from pynput.keyboard import Key, Controller
from layout import Layout
import docx
import os
from shutil import copyfile
import xlrd
import openpyxl as op

Prevention = True ## 是否防误触 设置时注意一下
task_id = 3  ##0,1,2,3 做完一轮任务改一下

def docx_len(name):
    file = docx.Document(name)
    para_data = []
    all_tables = file.tables
    for table in all_tables:
        for j in range(len(table.columns)):
            for i in range(len(table.rows)):
                para_single = table.cell(i,j).text
                para_data.extend(para_single)
    return len(para_data)

def record_screen(is_end, file_name):
    FPS = 20

    screenshot = ImageGrab.grab()
    H, W = screenshot.size
    video = cv2.VideoWriter('data/' + file_name + '.avi', cv2.VideoWriter_fourcc(*'XVID'), FPS, (H, W))
    timestamps = []

    while is_end.qsize() == 0:
        screenshot = ImageGrab.grab()
        frame = cv2.cvtColor(np.array(screenshot), cv2.COLOR_RGB2BGR)
        video.write(frame)
        timestamps.append(time.perf_counter())
    
    pickle.dump(timestamps, open('data/' + file_name + '.timestamp', 'wb'))
    video.release()
    print('Video released.', time.perf_counter())

def record_board(is_end, file_name):
    board = Board()
    history = History()
    [scalar, clf] = pickle.load(open('model/tap.model', 'rb'))
    labels = [None for i in range(20)]
    controller = Controller()
    qwerty = Layout()
    work_book = xlrd.open_workbook('data.xlsx')
    sheet_1 = work_book.sheet_by_index(0)
    row_sum = sheet_1.nrows
    wb = op.load_workbook("data.xlsx")
    sh=wb["Sheet1"]
    

    pygame.mixer.init(22050, -16, 2, 64)
    pygame.init()
    sound = pygame.mixer.Sound("sound/type.wav")
    sound.set_volume(1.0)
   
    os.startfile('task'+str(task_id)+'.docx')
    Ori_Words = docx_len('task'+str(task_id)+'.docx')
    print(Ori_Words)
    t1 = 0
    while is_end.qsize() == 0:
        if time.perf_counter() - t1 > 0.04 :
            print(time.perf_counter() - t1)
        t1 = time.perf_counter()
        frame = board.getNewFrame()
        history.updateFrame(frame)

        key_contacts = history.getKeyContact(frame)
        key_contacts = [contact for contact in key_contacts if labels[contact.id] == None]
        if len(key_contacts) > 0:
            if Prevention == True:
                features = [contact.feature for contact in key_contacts]
                features = scalar.transform(features)
                preds = clf.predict(features)
                j = 0
                for contact in key_contacts:
                    if preds[j]:
                        sound.play()
                        key = qwerty.decode(contact.x, contact.y)
                        if key == None :
                            continue
                        controller.press(key)
                        labels[contact.id] = key
                    j += 1
            else:
                j = 0
                for contact in key_contacts:
                    sound.play()
                    key = qwerty.decode(contact.x, contact.y)
                    if key == None :
                        continue
                    controller.press(key)
                    labels[contact.id] = key
                    j += 1

        for contact in frame.contacts:
            if contact.state == 1:
                labels[contact.id] = None
            if contact.state == 3:
                if labels[contact.id] != None:
                    controller.release(labels[contact.id])
            contact.label = (labels[contact.id] != None)
            contact.label = 0
        #frame.output()

    cv2.destroyAllWindows()
    board.stop()
    print('Board compressing.', time.perf_counter())
    compress_pickle.dump(board.frames, 'data/' + file_name + '.gz')
    print('Board released.', time.perf_counter())

    controller.press(Key.ctrl)
    controller.press('s')
    time.sleep(0.5)
    controller.release(Key.ctrl)
    controller.release('s')
    time.sleep(1)
    controller.press(Key.ctrl)
    controller.press('w')
    time.sleep(0.5)
    controller.release(Key.ctrl)
    controller.release('w')
    d = docx.Document('task'+str(task_id)+'.docx')
    d.save('data/' + file_name + '.docx')
    #print(docx_len('data/' + file_name + '.docx'))
    words = docx_len('data/' + file_name + '.docx') - Ori_Words
    Del_Words = qwerty.del_time
    Actual_time = qwerty.end_time - qwerty.start_time
    speed = words / Actual_time * 60
    accuracy = words / (Del_Words + words) * 100

    sh.cell(row_sum + 1,1,file_name)
    sh.cell(row_sum + 1,2,words)
    sh.cell(row_sum + 1,3,speed)
    sh.cell(row_sum + 1,4,Del_Words)
    sh.cell(row_sum + 1,5,accuracy)
    sh.cell(row_sum + 1,6,qwerty.back_times)
    print('Typing Words.', words)
    print('Typing Speed.', speed, 'WPM')
    print('Typing Accuracy', accuracy, '%')
    print('Back times', qwerty.back_times)
    
    wb.save("data.xlsx")


if __name__ == "__main__":
    file_name = DataManager().getFileName()
    is_end_1 = multiprocessing.Queue()
    is_end_2 = multiprocessing.Queue()
    p1 = multiprocessing.Process(target=record_screen, args=(is_end_1, file_name, ))
    p2 = multiprocessing.Process(target=record_board, args=(is_end_2, file_name, ))
    p1.start()
    p2.start()
    
    keyboard = MyKeyboard()
    start_time = time.perf_counter()
    while True:
        if keyboard.is_pressed('Esc') or time.perf_counter() - start_time > 360:
            is_end_1.put(1)
            is_end_2.put(1)
            break
        time.sleep(0.05)
    
    p1.join()
    p2.join()
    copyfile('C:\\Users\\chenxz\\Touch Type\\task'+str(task_id)+'.docx','task'+str(task_id)+'.docx')