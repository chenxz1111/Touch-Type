import docx
from data_manager import DataManager
import xlrd
import openpyxl as op
import re


def ori_word(name, id):
    file = docx.Document(name)
    all_tables = file.tables
    table = all_tables[id]
    para_single = table.cell(1,0).text
    para_single = re.sub(r'[^\u4e00-\u9fa5]','',para_single)
    return para_single

def act_word(name, id):
    file = docx.Document(name)
    all_tables = file.tables
    table = all_tables[id]
    para_single = table.cell(1,1).text
    para_single = re.sub(r'[^\u4e00-\u9fa5]','',para_single)
    return para_single


def med(str1, str2):
    assert(isinstance(str1, str) and isinstance(str2, str))
    m, n = len(str1), len(str2)
    matrix = [0 for i in range(2 * (n + 1))]
    for j in range(n + 1):
        matrix[j] = j
    for i in range(1, m + 1):
        for j in range(n + 1):
            if j == 0:
                matrix[j + n + 1] = i
                continue
            d = 0 if str1[i - 1] == str2[j - 1] else 1
            matrix[j + n + 1] = min(matrix[j + n] + 1, matrix[j] + 1, matrix[j - 1] + d)
        for j in range(n + 1):
            matrix[j] = matrix[j + n + 1]
    return matrix[n]

file_name = DataManager().getFileName()
name = 'data/' + file_name + '.docx'
id = int(file_name[-1]) - 1
wrong = med(ori_word(name, id).strip(), act_word(name, id).strip())
print(wrong)

print(len(ori_word(name ,id)))
print(act_word(name ,id))
